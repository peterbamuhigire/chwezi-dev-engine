#!/usr/bin/env python3
"""Build the reference.docx template used by the Pandoc production pipeline.

This script creates ``templates/reference.docx`` next to the skill. Pandoc maps
Markdown elements to the named Word styles defined here when invoked with
``--reference-doc=templates/reference.docx``. Editing the constants below is the
single supported way to change global document styling; never apply direct
formatting to a generated document.

Usage:
    python scripts/create-reference-docx.py [output_path]

Default output: ../templates/reference.docx relative to this script.

Dependencies:
    pip install python-docx
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


# --- Brand palette ---------------------------------------------------------
NAVY = RGBColor(0x1F, 0x38, 0x64)   # Title, Heading 1, table header fill
STEEL = RGBColor(0x2E, 0x5D, 0x8A)  # Heading 2
ACCENT = RGBColor(0x44, 0x72, 0xC4)  # Heading 3
BODY_INK = RGBColor(0x26, 0x26, 0x26)
CODE_INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x59, 0x59, 0x59)   # Caption, header, footer

# --- Font stacks -----------------------------------------------------------
FONT_HEADING = "Calibri Light"
FONT_BODY = "Calibri"
FONT_CODE = "Consolas"


def set_run_fonts(font_obj, name: str) -> None:
    """Set Latin and East Asian font names so Word does not substitute."""
    font_obj.name = name
    rpr = font_obj.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), name)


def style_font(style, *, name: str, size: float, bold: bool, colour: RGBColor) -> None:
    font = style.font
    set_run_fonts(font, name)
    font.size = Pt(size)
    font.bold = bold
    font.color.rgb = colour


def spacing(style, *, before: float, after: float, line: float | None,
            line_rule=WD_LINE_SPACING.SINGLE) -> None:
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line
        pf.line_spacing_rule = line_rule


def add_left_border(style, *, size_eighths: int, colour_hex: str) -> None:
    """Add a coloured left border bar to a paragraph style (the H1 anchor)."""
    ppr = style.element.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_eighths))  # in eighths of a point
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), colour_hex)
    pbdr.append(left)


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    # Normal / body
    normal = styles["Normal"]
    style_font(normal, name=FONT_BODY, size=11, bold=False, colour=BODY_INK)
    spacing(normal, before=0, after=6, line=1.15,
            line_rule=WD_LINE_SPACING.MULTIPLE)
    npf = normal.paragraph_format
    npf.keep_together = True       # never split a paragraph mid-page
    npf.widow_control = True       # no widows or orphans

    # Title
    title = styles["Title"]
    style_font(title, name=FONT_HEADING, size=28, bold=True, colour=NAVY)
    spacing(title, before=0, after=12, line=None)

    # Subtitle
    if "Subtitle" in [s.name for s in styles]:
        subtitle = styles["Subtitle"]
        style_font(subtitle, name=FONT_HEADING, size=14, bold=False, colour=STEEL)
        spacing(subtitle, before=0, after=18, line=None)

    # Heading 1 - new page, navy left border anchor
    h1 = styles["Heading 1"]
    style_font(h1, name=FONT_HEADING, size=16, bold=True, colour=NAVY)
    spacing(h1, before=20, after=6, line=None)
    h1.paragraph_format.page_break_before = True
    h1.paragraph_format.keep_with_next = True
    add_left_border(h1, size_eighths=36, colour_hex="1F3864")  # 4.5pt bar

    # Heading 2
    h2 = styles["Heading 2"]
    style_font(h2, name=FONT_HEADING, size=13, bold=True, colour=STEEL)
    spacing(h2, before=14, after=4, line=None)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True

    # Heading 3
    h3 = styles["Heading 3"]
    style_font(h3, name=FONT_BODY, size=11, bold=True, colour=ACCENT)
    spacing(h3, before=10, after=3, line=None)
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.keep_together = True

    # Heading 4
    if "Heading 4" in [s.name for s in styles]:
        h4 = styles["Heading 4"]
        style_font(h4, name=FONT_BODY, size=11, bold=True, colour=BODY_INK)
        spacing(h4, before=8, after=3, line=None)
        h4.paragraph_format.keep_with_next = True
        h4.paragraph_format.keep_together = True

    # Code / verbatim blocks
    for code_name in ("Source Code", "Verbatim", "Macro Text"):
        if code_name in [s.name for s in styles]:
            cs = styles[code_name]
            style_font(cs, name=FONT_CODE, size=9.5, bold=False, colour=CODE_INK)

    # Caption
    if "Caption" in [s.name for s in styles]:
        cap = styles["Caption"]
        style_font(cap, name=FONT_BODY, size=9, bold=False, colour=MUTED)
        spacing(cap, before=2, after=8, line=None)


def configure_page_and_footer(doc: Document) -> None:
    section = doc.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    # Cover page carries no header/footer.
    section.different_first_page_header_footer = True

    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.text = "Page "
    run = para.runs[0]
    set_run_fonts(run.font, FONT_BODY)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    _add_page_field(para)


def _add_page_field(paragraph) -> None:
    """Insert a live PAGE field so footers show the current page number."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_fonts(run.font, FONT_BODY)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def build(output_path: Path) -> None:
    doc = Document()
    configure_styles(doc)
    configure_page_and_footer(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Wrote reference template: {output_path}")


def main(argv: list[str]) -> int:
    if len(argv) > 1:
        output = Path(argv[1])
    else:
        output = Path(__file__).resolve().parent.parent / "templates" / "reference.docx"
    try:
        build(output)
    except ModuleNotFoundError:
        print("python-docx is required. Install it with: pip install python-docx",
              file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
